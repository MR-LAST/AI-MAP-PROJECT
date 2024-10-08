from docx import Document

# Create a new Document
doc = Document()

# Title Slide
doc.add_heading('Marketing Proposal for Try Rooftop', 0)
doc.add_paragraph('Boosting Visibility, Engagement, and Foot Traffic through Tailored Marketing Campaigns')
doc.add_paragraph('Presented by: Trident Marketing')
doc.add_paragraph('Date: [Insert Date]')

# Introduction to Trident Marketing
doc.add_heading('Introduction to Trident Marketing', level=1)
doc.add_paragraph(
    "About Us:\n"
    "Trident Marketing specializes in creating innovative, results-driven marketing campaigns. "
    "We have a proven track record in helping businesses increase visibility and customer engagement."
)
doc.add_paragraph(
    "Our Objective: To position Try Rooftop as a premier destination for dining, entertainment, "
    "and social gatherings in [City Name]."
)

# Proposal Overview
doc.add_heading('Proposal Overview', level=1)
doc.add_paragraph(
    "Goal: To draw more people to Try Rooftop by utilizing a strategic mix of marketing, ads, "
    "and promotions."
)
doc.add_paragraph(
    "Our Approach: Focused on increasing foot traffic, driving online engagement, and creating "
    "lasting customer loyalty."
)

# Save the document
file_path = "/mnt/data/Marketing_Proposal_Try_Rooftop.docx"
doc.save(file_path)

file_path

